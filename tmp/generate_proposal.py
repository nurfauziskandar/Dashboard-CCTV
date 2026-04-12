#!/usr/bin/env python3
"""
Generator Dokumen Seminar Proposal Tugas Akhir
Politeknik Penerbangan Indonesia Curug (PPIC)
Format: BAB I – BAB III

Output: ../Proposal_Seminar_TA_Nurfauzi_Iskandar.docx
"""

import os
import sys
import tempfile

# ── pastikan tmp/ ada di path ──────────────────────────────────────────────
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
from flowchart_builder import (build_linear_flowchart,
                                build_branching_flowchart,
                                build_block_diagram)

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Identitas ─────────────────────────────────────────────────────────────────
MAHASISWA_NAMA = "Nurfauzi Iskandar"
MAHASISWA_NIM  = "xxxx-xx-xxx"            # ← ganti NIM asli
PRODI          = "D4 Teknik Listrik Bandar Udara"
JURUSAN        = "Teknologi Penerbangan"
INSTITUSI      = "POLITEKNIK PENERBANGAN INDONESIA CURUG"
KOTA           = "Curug, Tangerang"
TAHUN          = "2026"
BULAN          = "Maret 2026"
DOSEN_1        = "_________________, S.T., M.T."   # ← ganti pembimbing I
DOSEN_2        = "_________________, S.T., M.T."   # ← ganti pembimbing II

JUDUL = (
    "RANCANG BANGUN DASHBOARD MONITORING KAMERA CCTV BERBASIS WEB\n"
    "MENGGUNAKAN PROTOKOL ONVIF PADA SISTEM KEAMANAN\n"
    "BANDAR UDARA SOEKARNO-HATTA"
)
JUDUL_SATU_BARIS = (
    "RANCANG BANGUN DASHBOARD MONITORING KAMERA CCTV BERBASIS WEB "
    "MENGGUNAKAN PROTOKOL ONVIF PADA SISTEM KEAMANAN "
    "BANDAR UDARA SOEKARNO-HATTA"
)

OUTPUT_PATH = os.path.join(
    os.path.dirname(os.path.dirname(os.path.abspath(__file__))),
    "Proposal_Seminar_TA_Nurfauzi_Iskandar.docx"
)

# Folder output PNG flowchart — disimpan permanen di tmp/
_TMP_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), "flowcharts")
os.makedirs(_TMP_DIR, exist_ok=True)


# ════════════════════════════════════════════════════════════════════════════
#  FLOWCHART PNG GENERATORS
# ════════════════════════════════════════════════════════════════════════════

def fc_path(name):
    return os.path.join(_TMP_DIR, f"{name}.png")


def generate_all_flowcharts():
    print("  Membuat diagram PNG ...")

    # Gambar 1.1 — Diagram Blok Sistem (baru)
    build_block_diagram(fc_path("block_diagram"))

    # Gambar 2.1 — Kerangka Berpikir
    build_linear_flowchart(
        "Kerangka Berpikir Penelitian",
        [
            ("start_end", "MULAI"),
            ("process",   "Identifikasi Masalah"),
            ("process",   "Kajian Pustaka"),
            ("process",   "Usulan Solusi"),
            ("process",   "Perancangan Sistem"),
            ("process",   "Implementasi"),
            ("decision",  "Sistem sesuai\nkebutuhan?"),
            ("process",   "Pengujian &\nEvaluasi"),
            ("start_end", "SELESAI"),
        ],
        fc_path("kerangka_berpikir"),
    )

    # Gambar 3.1 — Alur Perancangan Dashboard
    build_linear_flowchart(
        "Flowchart Alur Perancangan Dashboard",
        [
            ("start_end", "MULAI"),
            ("process",   "Analisis Kebutuhan Sistem"),
            ("process",   "Perancangan Arsitektur Sistem"),
            ("process",   "Perancangan Basis Data"),
            ("process",   "Perancangan Antarmuka Pengguna"),
            ("process",   "Implementasi Sistem"),
            ("process",   "Pengujian Fungsional"),
            ("decision",  "Sistem sesuai\nrancangan?"),
            ("process",   "Dokumentasi Sistem"),
            ("start_end", "SELESAI"),
        ],
        fc_path("alur_perancangan"),
    )

    # Gambar 3.2 — Flowchart Pemantauan Kamera
    build_linear_flowchart(
        "Flowchart Pemantauan Status Kamera",
        [
            ("start_end", "MULAI"),
            ("process",   "Ambil data kamera\ndari sistem"),
            ("decision",  "Kamera\ntersedia?"),
            ("process",   "Periksa status\nkamera"),
            ("decision",  "Kamera\naktif?"),
            ("process",   "Tandai kamera\nAktif"),
            ("process",   "Tandai kamera\nTidak Aktif"),
            ("process",   "Simpan hasil\npemantauan"),
            ("start_end", "SELESAI"),
        ],
        fc_path("polling_kamera"),
    )

    # Gambar 3.2 — Flowchart Pemantauan Server Storage (branching)
    build_branching_flowchart(
        "Flowchart Pemantauan Server Storage",
        decision_text="Jenis\nServer?",
        steps_left=[
            ("process", "Ambil data\nsuhu server"),
            ("process", "Ambil data\nkondisi HDD"),
            ("process", "Ambil data\ndaya listrik"),
        ],
        label_left="VX Storage",
        steps_right=[
            ("process", "Ambil info\nsistem"),
            ("process", "Ambil data\npenyimpanan"),
        ],
        label_right="Endura",
        merge_step=("process", "Simpan hasil\npemantauan server"),
        output_path=fc_path("polling_server"),
    )

    # Gambar 3.3 — Diagram Alir Penelitian
    build_linear_flowchart(
        "Diagram Alir Penelitian",
        [
            ("start_end", "MULAI"),
            ("process",   "Identifikasi Masalah\n& Studi Literatur"),
            ("process",   "Perumusan Masalah\n& Tujuan Penelitian"),
            ("process",   "Perancangan Sistem"),
            ("process",   "Implementasi\nSistem"),
            ("decision",  "Sistem berjalan\ndengan baik?"),
            ("process",   "Perbaikan Sistem"),
            ("process",   "Pengujian &\nValidasi"),
            ("decision",  "Pengujian\nberhasil?"),
            ("process",   "Analisis Hasil\n& Pembahasan"),
            ("process",   "Penulisan\nLaporan"),
            ("start_end", "SELESAI"),
        ],
        fc_path("alir_penelitian"),
    )

    print(f"    → 5 diagram PNG disimpan di {_TMP_DIR}")


# ════════════════════════════════════════════════════════════════════════════
#  HELPER: Format Paragraf
# ════════════════════════════════════════════════════════════════════════════

def set_font(run, size=12, bold=False, italic=False, color=None):
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)
    rPr = run._r.get_or_add_rPr()
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), "Times New Roman")
    rFonts.set(qn("w:hAnsi"), "Times New Roman")
    rPr.insert(0, rFonts)


def set_para_fmt(para, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
                 space_before=0, space_after=6,
                 first_indent=None, left_indent=None):
    pf = para.paragraph_format
    pf.alignment = align
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = 1.5
    if first_indent is not None:
        pf.first_line_indent = Cm(first_indent)
    if left_indent is not None:
        pf.left_indent = Cm(left_indent)


def body(doc, text, indent=True, space_after=6, bold=False, italic=False):
    para = doc.add_paragraph()
    set_para_fmt(para, space_after=space_after,
                 first_indent=1.25 if indent else 0)
    run = para.add_run(text)
    set_font(run, size=12, bold=bold, italic=italic)
    return para


def chapter_title(doc, text):
    para = doc.add_paragraph()
    para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after = Pt(12)
    para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    para.paragraph_format.line_spacing = 1.5
    run = para.add_run(text.upper())
    set_font(run, size=14, bold=True)


def section_h(doc, number, text):
    para = doc.add_paragraph()
    set_para_fmt(para, align=WD_ALIGN_PARAGRAPH.LEFT,
                 space_before=12, space_after=6, first_indent=0)
    run = para.add_run(f"{number}  {text}")
    set_font(run, size=12, bold=True)


def subsection_h(doc, number, text):
    para = doc.add_paragraph()
    set_para_fmt(para, align=WD_ALIGN_PARAGRAPH.LEFT,
                 space_before=6, space_after=6, first_indent=0)
    run = para.add_run(f"{number}  {text}")
    set_font(run, size=12, bold=True)


def add_figure(doc, png_path, caption_text, width_cm=12.0):
    """Sisipkan PNG dan caption di bawahnya."""
    para_img = doc.add_paragraph()
    para_img.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para_img.paragraph_format.space_before = Pt(6)
    para_img.paragraph_format.space_after = Pt(0)
    run = para_img.add_run()
    run.add_picture(png_path, width=Cm(width_cm))

    para_cap = doc.add_paragraph()
    para_cap.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para_cap.paragraph_format.space_before = Pt(2)
    para_cap.paragraph_format.space_after = Pt(10)
    para_cap.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    r = para_cap.add_run(caption_text)
    set_font(r, size=11, bold=False)


def add_table(doc, headers, rows, caption_text):
    """Tabel dengan caption di atas."""
    # caption di atas
    para_cap = doc.add_paragraph()
    para_cap.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    para_cap.paragraph_format.space_before = Pt(8)
    para_cap.paragraph_format.space_after = Pt(3)
    para_cap.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    r = para_cap.add_run(caption_text)
    set_font(r, size=11, bold=False)

    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    hdr_row = tbl.rows[0]
    for i, h in enumerate(headers):
        cell = hdr_row.cells[i]
        cell.text = h
        for run in cell.paragraphs[0].runs:
            set_font(run, size=11, bold=True)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "D9D9D9")
        tcPr.append(shd)

    for r_i, row_data in enumerate(rows):
        row = tbl.rows[r_i + 1]
        for c_i, val in enumerate(row_data):
            cell = row.cells[c_i]
            cell.text = str(val)
            for run in cell.paragraphs[0].runs:
                set_font(run, size=11)

    doc.add_paragraph()


def numbered_list(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.left_indent = Cm(1.25)
        run = p.add_run(item)
        set_font(run, size=12)


def set_margins(section):
    section.left_margin   = Cm(4)
    section.right_margin  = Cm(3)
    section.top_margin    = Cm(3)
    section.bottom_margin = Cm(3)


# ════════════════════════════════════════════════════════════════════════════
#  HALAMAN JUDUL
# ════════════════════════════════════════════════════════════════════════════
def build_cover(doc):
    def ctr(text, size=12, bold=False, single=True):
        p = doc.add_paragraph()
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(4)
        if single:
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        else:
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
            p.paragraph_format.line_spacing = 1.5
        r = p.add_run(text)
        set_font(r, size=size, bold=bold)
        return p

    ctr("PROPOSAL TUGAS AKHIR", size=14, bold=True)
    doc.add_paragraph()

    p_judul = doc.add_paragraph()
    p_judul.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_judul.paragraph_format.space_before = Pt(6)
    p_judul.paragraph_format.space_after = Pt(6)
    p_judul.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p_judul.paragraph_format.line_spacing = 1.5
    r = p_judul.add_run(JUDUL)
    set_font(r, size=14, bold=True)

    doc.add_paragraph()
    doc.add_paragraph()

    p_logo = doc.add_paragraph()
    p_logo.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_logo.paragraph_format.space_after = Pt(8)
    r_logo = p_logo.add_run("[LOGO PPIC CURUG]")
    set_font(r_logo, size=12, italic=True, color=(128, 128, 128))

    doc.add_paragraph()
    ctr("Disusun oleh:", size=12)
    ctr(MAHASISWA_NAMA.upper(), size=12, bold=True)
    ctr(f"NIM: {MAHASISWA_NIM}", size=12, bold=True)
    doc.add_paragraph()
    ctr(f"Program Studi {PRODI}", size=12)
    ctr(f"Jurusan {JURUSAN}", size=12)
    ctr(INSTITUSI, size=12, bold=True)
    ctr(BULAN, size=12)


# ════════════════════════════════════════════════════════════════════════════
#  HALAMAN PERSETUJUAN
# ════════════════════════════════════════════════════════════════════════════
def build_approval(doc):
    doc.add_page_break()

    p = doc.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(18)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    r = p.add_run("HALAMAN PERSETUJUAN")
    set_font(r, size=14, bold=True)

    body(doc, "Proposal Tugas Akhir dengan judul:", indent=False)

    p_j = doc.add_paragraph()
    p_j.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_j.paragraph_format.space_before = Pt(6)
    p_j.paragraph_format.space_after = Pt(12)
    p_j.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p_j.paragraph_format.line_spacing = 1.5
    r = p_j.add_run(f'"{JUDUL_SATU_BARIS}"')
    set_font(r, size=12, bold=True)

    body(doc,
         f"Disusun oleh {MAHASISWA_NAMA} (NIM: {MAHASISWA_NIM}), "
         f"mahasiswa Program Studi {PRODI}, Jurusan {JURUSAN}, "
         f"{INSTITUSI}, telah disetujui dan disahkan pada tanggal ________________.",
         indent=False, space_after=18)

    # Tanda tangan pembimbing
    tbl = doc.add_table(rows=4, cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (jabatan, nama) in enumerate([("Pembimbing I", DOSEN_1),
                                          ("Pembimbing II", DOSEN_2)]):
        tbl.rows[0].cells[i].text = jabatan
        tbl.rows[0].cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in tbl.rows[0].cells[i].paragraphs[0].runs:
            set_font(run, size=12)
        tbl.rows[2].cells[i].text = "\n\n\n"
        tbl.rows[3].cells[i].text = nama
        tbl.rows[3].cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in tbl.rows[3].cells[i].paragraphs[0].runs:
            set_font(run, size=12, bold=True)

    doc.add_paragraph()
    doc.add_paragraph()

    p_kps = doc.add_paragraph()
    p_kps.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p_kps.add_run(f"Mengetahui,\nKetua Program Studi {PRODI}")
    set_font(r, size=12)

    doc.add_paragraph("\n\n\n").paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p_kap = doc.add_paragraph()
    p_kap.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p_kap.add_run("_________________________\nNIP. _____________________")
    set_font(r, size=12)


# ════════════════════════════════════════════════════════════════════════════
#  DAFTAR ISI
# ════════════════════════════════════════════════════════════════════════════
def build_daftar_isi(doc):
    doc.add_page_break()
    p = doc.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(18)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    r = p.add_run("DAFTAR ISI")
    set_font(r, size=14, bold=True)

    entries = [
        ("HALAMAN JUDUL", "i"),
        ("HALAMAN PERSETUJUAN", "ii"),
        ("DAFTAR ISI", "iii"),
        ("DAFTAR GAMBAR", "iv"),
        ("DAFTAR TABEL", "v"),
        ("BAB I  PENDAHULUAN", "1"),
        ("    1.1  Latar Belakang Masalah", "1"),
        ("    1.2  Rumusan Masalah", "5"),
        ("    1.3  Batasan Masalah", "6"),
        ("    1.4  Tujuan Penelitian", "6"),
        ("    1.5  Manfaat Penelitian", "7"),
        ("    1.6  Sistematika Penulisan", "8"),
        ("BAB II  TINJAUAN PUSTAKA", "9"),
        ("    2.1  Sistem CCTV pada Bandar Udara", "9"),
        ("    2.2  Protokol ONVIF", "11"),
        ("    2.3  Kamera dan Perangkat Pelco", "13"),
        ("    2.4  Flask Web Framework", "15"),
        ("    2.5  Teknologi Frontend Web", "17"),
        ("    2.6  Monitoring Server Storage", "19"),
        ("    2.7  Sistem Basis Data SQLite", "21"),
        ("    2.8  Kerangka Berpikir", "22"),
        ("BAB III  METODOLOGI PENELITIAN", "23"),
        ("    3.1  Jenis Penelitian", "23"),
        ("    3.2  Waktu dan Tempat Penelitian", "24"),
        ("    3.3  Teknik Pengumpulan Data", "24"),
        ("    3.4  Perancangan Sistem", "25"),
        ("    3.5  Diagram Alir Penelitian", "31"),
        ("    3.6  Jadwal Kegiatan Penelitian", "32"),
        ("DAFTAR PUSTAKA", "33"),
    ]

    for label, page in entries:
        p_e = doc.add_paragraph()
        p_e.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p_e.paragraph_format.space_before = Pt(0)
        p_e.paragraph_format.space_after = Pt(3)
        p_e.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        p_e.paragraph_format.tab_stops.add_tab_stop(
            Cm(13.5), WD_ALIGN_PARAGRAPH.RIGHT)
        bold = (label.startswith("BAB") or
                label in ("HALAMAN JUDUL", "HALAMAN PERSETUJUAN",
                          "DAFTAR ISI", "DAFTAR GAMBAR",
                          "DAFTAR TABEL", "DAFTAR PUSTAKA"))
        r = p_e.add_run(f"{label}\t{page}")
        set_font(r, size=12, bold=bold)


# ════════════════════════════════════════════════════════════════════════════
#  BAB I – PENDAHULUAN
# ════════════════════════════════════════════════════════════════════════════
def build_bab1(doc):
    doc.add_page_break()
    chapter_title(doc, "BAB I\nPENDAHULUAN")

    section_h(doc, "1.1", "Latar Belakang Masalah")

    body(doc,
         "Bandar Udara Soekarno-Hatta merupakan salah satu bandar udara tersibuk "
         "di Asia Tenggara dengan pergerakan penumpang mencapai lebih dari 60 juta "
         "orang per tahun. Sebagai infrastruktur kritikal dalam industri penerbangan "
         "nasional, keamanan bandar udara menjadi aspek yang tidak dapat dikesampingkan. "
         "Salah satu komponen utama dalam sistem keamanan bandar udara adalah penggunaan "
         "kamera CCTV (Closed-Circuit Television) yang terpasang di berbagai titik "
         "strategis, mulai dari terminal keberangkatan dan kedatangan, apron, taxiway, "
         "hingga area parkir dan perimeter.")

    body(doc,
         "Sistem CCTV yang digunakan pada Bandar Udara Soekarno-Hatta didominasi oleh "
         "kamera dari merek Pelco, yang merupakan salah satu produsen kamera pengawas "
         "profesional terkemuka di dunia. Kamera-kamera Pelco tersebut mencakup berbagai "
         "seri, antara lain Sarix IMP, Sarix IME, Spectra Pro, dan Optera, yang seluruhnya "
         "mendukung standar ONVIF (Open Network Video Interface Forum) Profile S dan T. "
         "Standar ONVIF memungkinkan interoperabilitas antara perangkat kamera dari "
         "berbagai produsen dengan sistem manajemen video (VMS) yang ada.")

    body(doc,
         "Namun, pengelolaan ratusan kamera CCTV secara bersamaan menimbulkan tantangan "
         "tersendiri bagi tim teknisi keamanan bandar udara. Permasalahan yang kerap "
         "dihadapi antara lain: (1) sulitnya memantau status operasional setiap kamera "
         "secara real-time, (2) tidak tersedianya informasi terpusat mengenai kondisi "
         "kamera aktif dan tidak aktif, (3) keterbatasan visibilitas terhadap lokasi "
         "geografis setiap kamera, dan (4) tidak adanya sistem monitoring otomatis untuk "
         "server penyimpanan rekaman yang digunakan oleh sistem NVR (Network Video Recorder).")

    body(doc,
         "Server penyimpanan rekaman CCTV menggunakan sistem Pelco VX Storage E-Series, "
         "yang secara fisik merupakan server Dell PowerEdge R740xd dengan iDRAC "
         "(Integrated Dell Remote Access Controller) generasi ke-9 built-in. Selain itu, "
         "terdapat pula sistem Pelco Endura NSM5200 (Network Storage Manager) sebagai "
         "pengelola penyimpanan berbasis jaringan. Kondisi kesehatan hard disk drive "
         "(HDD) pada server-server ini sangat krusial, karena kegagalan HDD dapat "
         "mengakibatkan hilangnya rekaman video yang telah tersimpan. Monitoring kondisi "
         "HDD selama ini dilakukan secara manual dan sporadis.")

    body(doc,
         "Untuk mengatasi permasalahan tersebut, diperlukan sebuah sistem dashboard "
         "monitoring terpadu yang mampu menampilkan status seluruh kamera CCTV dan "
         "server storage secara real-time melalui antarmuka web. Penelitian ini mengusulkan "
         "rancang bangun aplikasi Dashboard Monitoring CCTV berbasis web menggunakan "
         "framework Flask (Python) dengan integrasi protokol ONVIF melalui library "
         "onvif-zeep, peta interaktif Leaflet.js, monitoring iDRAC Redfish REST API, "
         "dan SNMP untuk Endura NSM5200. Gambaran umum arsitektur sistem yang diusulkan "
         "dapat dilihat pada Gambar 1.1. Dengan demikian, diharapkan sistem ini dapat "
         "meningkatkan efisiensi pemeliharaan dan respons terhadap gangguan sistem "
         "keamanan di Bandar Udara Soekarno-Hatta.")

    add_figure(doc, fc_path("block_diagram"),
               "Gambar 1.1  Diagram Blok Sistem Dashboard Monitoring CCTV",
               width_cm=15.5)

    section_h(doc, "1.2", "Rumusan Masalah")
    body(doc, "Berdasarkan latar belakang di atas, rumusan masalah dalam penelitian ini adalah:")
    numbered_list(doc, [
        "Bagaimana merancang dan membangun sistem dashboard monitoring kamera CCTV "
        "berbasis web yang menampilkan status kamera secara real-time menggunakan protokol ONVIF?",
        "Bagaimana mengintegrasikan monitoring hardware server storage Pelco VX Storage "
        "melalui Dell iDRAC Redfish API dan Pelco Endura NSM5200 melalui SNMP ke dalam "
        "satu antarmuka dashboard terpadu?",
        "Bagaimana menampilkan informasi lokasi geografis kamera CCTV pada peta interaktif "
        "untuk memudahkan identifikasi posisi kamera di area bandar udara?",
        "Bagaimana merancang sistem background polling otomatis yang memperbarui status "
        "kamera dan server storage secara berkala tanpa mengganggu kinerja aplikasi?",
    ])

    section_h(doc, "1.3", "Batasan Masalah")
    body(doc, "Batasan masalah dalam penelitian ini adalah sebagai berikut:")
    numbered_list(doc, [
        "Sistem monitoring hanya mendukung kamera CCTV yang kompatibel dengan standar "
        "ONVIF Profile S dan T, khususnya produk Pelco seri Sarix, Spectra, dan Optera.",
        "Monitoring server storage terbatas pada dua jenis perangkat: (a) Pelco VX Storage "
        "melalui Dell iDRAC 9 Redfish API, dan (b) Pelco Endura NSM5200 melalui SNMP v2c/v3.",
        "Aplikasi dikembangkan menggunakan Python 3 dengan Flask 3 dan basis data SQLite, "
        "tanpa server database eksternal.",
        "Streaming video kamera ditampilkan dalam format MJPEG melalui konversi RTSP "
        "menggunakan library OpenCV.",
        "Aplikasi tidak mencakup fitur autentikasi multi-user dan sistem perekaman video.",
        "Pengujian sistem dilakukan dalam lingkungan jaringan lokal (LAN) menggunakan "
        "data simulasi (demo mode).",
    ])

    section_h(doc, "1.4", "Tujuan Penelitian")
    body(doc, "Tujuan yang ingin dicapai dalam penelitian ini adalah:")
    numbered_list(doc, [
        "Merancang dan membangun aplikasi dashboard monitoring kamera CCTV berbasis web "
        "yang mengintegrasikan protokol ONVIF untuk pemantauan status kamera secara real-time.",
        "Mengimplementasikan integrasi monitoring hardware server storage Pelco melalui "
        "Dell iDRAC Redfish REST API dan SNMP ke dalam antarmuka dashboard terpadu.",
        "Menampilkan persebaran lokasi kamera CCTV secara visual pada peta interaktif "
        "berbasis Leaflet.js dengan OpenStreetMap.",
        "Membangun sistem background polling otomatis menggunakan Flask-APScheduler "
        "untuk pembaruan status kamera dan server storage secara periodik.",
    ])

    section_h(doc, "1.5", "Manfaat Penelitian")
    body(doc, "Manfaat yang diharapkan dari penelitian ini meliputi:")

    p_a = doc.add_paragraph()
    set_para_fmt(p_a, align=WD_ALIGN_PARAGRAPH.LEFT, first_indent=0)
    r = p_a.add_run("a.  Manfaat Teoritis")
    set_font(r, size=12, bold=True)
    numbered_list(doc, [
        "Memberikan kontribusi pengembangan ilmu pengetahuan di bidang sistem informasi "
        "berbasis web dan integrasi protokol jaringan ONVIF.",
        "Menjadi referensi bagi penelitian selanjutnya tentang monitoring sistem keamanan "
        "berbasis protokol standar terbuka.",
    ])

    p_b = doc.add_paragraph()
    set_para_fmt(p_b, align=WD_ALIGN_PARAGRAPH.LEFT, first_indent=0)
    r = p_b.add_run("b.  Manfaat Praktis")
    set_font(r, size=12, bold=True)
    numbered_list(doc, [
        "Memudahkan tim teknisi bandar udara dalam memantau status operasional seluruh "
        "kamera CCTV dan server storage melalui satu antarmuka terpadu.",
        "Mempercepat identifikasi dan penanganan kamera tidak aktif atau server storage "
        "yang mengalami masalah hardware.",
        "Mengurangi risiko kehilangan rekaman video akibat kegagalan HDD yang tidak "
        "terdeteksi secara dini.",
        "Hasil penelitian dapat diimplementasikan di bandar udara lain dengan sistem "
        "Pelco dan infrastruktur sejenis.",
    ])

    section_h(doc, "1.6", "Sistematika Penulisan")
    body(doc, "Penulisan proposal tugas akhir ini disusun ke dalam tiga bab sebagai berikut:")

    for judul_bab, isi in [
        ("BAB I  PENDAHULUAN",
         "Menguraikan latar belakang masalah, rumusan masalah, batasan masalah, "
         "tujuan penelitian, manfaat penelitian, dan sistematika penulisan."),
        ("BAB II  TINJAUAN PUSTAKA",
         "Memaparkan landasan teori yang mencakup sistem CCTV pada bandar udara, "
         "protokol ONVIF, kamera dan perangkat Pelco, Flask web framework, teknologi "
         "frontend web, monitoring server storage, basis data SQLite, serta kerangka berpikir."),
        ("BAB III  METODOLOGI PENELITIAN",
         "Menjelaskan jenis penelitian, waktu dan tempat penelitian, teknik pengumpulan "
         "data, perancangan sistem beserta flowchart, diagram alir penelitian, dan "
         "jadwal kegiatan."),
    ]:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.first_line_indent = Cm(1.25)
        r1 = p.add_run(judul_bab + ":  ")
        set_font(r1, size=12, bold=True)
        r2 = p.add_run(isi)
        set_font(r2, size=12)


# ════════════════════════════════════════════════════════════════════════════
#  BAB II – TINJAUAN PUSTAKA
# ════════════════════════════════════════════════════════════════════════════
def build_bab2(doc):
    doc.add_page_break()
    chapter_title(doc, "BAB II\nTINJAUAN PUSTAKA")

    section_h(doc, "2.1", "Sistem CCTV pada Bandar Udara")
    body(doc,
         "CCTV (Closed-Circuit Television) adalah sistem pengawasan video yang menggunakan "
         "kamera untuk mengirimkan sinyal ke sekumpulan monitor atau perekam video dalam "
         "jaringan tertutup. Berdasarkan Annex 17 ICAO tentang keamanan penerbangan, "
         "seluruh bandar udara internasional diwajibkan memiliki sistem pemantauan keamanan "
         "yang memadai, termasuk kamera CCTV di area-area kritis, dengan kewajiban "
         "pengarsipan rekaman video minimal 30 hari untuk keperluan investigasi (ICAO, 2017).")
    body(doc,
         "Komponen utama sistem CCTV modern pada bandar udara terdiri dari: (1) kamera IP "
         "yang menangkap dan mengirimkan gambar melalui jaringan Ethernet, (2) NVR "
         "(Network Video Recorder) sebagai server penyimpanan rekaman, (3) VMS (Video "
         "Management System) sebagai perangkat lunak pengelola, (4) workstation operator "
         "untuk pemantauan langsung, dan (5) infrastruktur jaringan LAN/VLAN khusus "
         "untuk video surveillance.")
    body(doc,
         "Tantangan utama pengelolaan sistem CCTV skala besar adalah pemantauan kesehatan "
         "sistem secara real-time. Kamera yang tidak aktif akibat gangguan jaringan, daya, "
         "atau kerusakan fisik menimbulkan celah keamanan signifikan. Kapasitas penyimpanan "
         "dan kondisi fisik HDD pada server NVR perlu dipantau secara proaktif untuk "
         "mencegah kehilangan data rekaman.")

    section_h(doc, "2.2", "Protokol ONVIF")
    body(doc,
         "ONVIF (Open Network Video Interface Forum) adalah standar antarmuka terbuka yang "
         "dikembangkan oleh konsorsium industri — didirikan tahun 2008 oleh Axis, Bosch, "
         "dan Sony. Tujuan utamanya adalah memastikan interoperabilitas perangkat keamanan "
         "jaringan dari berbagai produsen (ONVIF, 2023). Komunikasi ONVIF menggunakan "
         "protokol SOAP berbasis XML yang ditransmisikan melalui HTTP pada port 80.")
    body(doc,
         "Profil utama ONVIF yang relevan: Profile S (Streaming) — mendukung konfigurasi "
         "video, streaming RTSP, dan PTZ; Profile T (Advanced Streaming) — menambahkan "
         "dukungan H.265, motion detection, dan metadata. Seluruh kamera Pelco yang "
         "digunakan mendukung minimum Profile S dan T (Pelco, 2022).")

    add_table(doc,
        ["Operasi ONVIF", "Protokol/Port", "Fungsi"],
        [
            ("GetDeviceInformation", "HTTP/TCP 80",
             "Mendapatkan informasi perangkat: pabrikan, model, firmware, serial number"),
            ("GetProfiles",          "HTTP/TCP 80",
             "Mendapatkan profil media yang didukung kamera"),
            ("GetStreamUri",         "HTTP/TCP 80",
             "Mendapatkan URL RTSP untuk video streaming"),
            ("GetSnapshotUri",       "HTTP/TCP 80",
             "Mendapatkan URL HTTP untuk pengambilan gambar JPEG"),
            ("WS-Discovery",         "UDP/3702 Multicast",
             "Penemuan otomatis perangkat ONVIF dalam jaringan lokal"),
        ],
        "Tabel 2.1  Operasi ONVIF yang digunakan dalam sistem dashboard")

    section_h(doc, "2.3", "Kamera dan Perangkat Pelco")
    body(doc,
         "Pelco adalah merek kamera pengawas profesional milik Motorola Solutions yang "
         "banyak digunakan di fasilitas kritis seperti bandar udara dan pelabuhan. Seri "
         "yang relevan: Sarix IMP/IME (mini dome, resolusi hingga 8MP), Spectra Enhanced "
         "PTZ (36x optical zoom, auto-tracking), dan Optera (panoramic 180°/360°). Semua "
         "seri menggunakan format RTSP standar: rtsp://<ip>:554/stream1 (primary) dan "
         "rtsp://<ip>:554/stream2 (substream).")
    body(doc,
         "Pelco VX Storage E-Series adalah NVR berbasis Dell PowerEdge R740xd dengan "
         "Dell iDRAC 9 built-in yang mendukung Redfish API — standar RESTful API berbasis "
         "JSON untuk manajemen hardware server secara remote (Dell Technologies, 2023). "
         "Pelco Endura NSM5200 adalah network storage manager yang mendukung monitoring "
         "melalui SNMP v2c/v3.")

    section_h(doc, "2.4", "Flask Web Framework")
    body(doc,
         "Flask adalah microframework web Python yang dikembangkan oleh Armin Ronacher. "
         "Disebut 'micro' karena tidak memaksa penggunaan library tertentu, memberikan "
         "fleksibilitas kepada pengembang (Grinberg, 2018). Flask menggunakan Werkzeug "
         "sebagai WSGI toolkit dan Jinja2 sebagai template engine.")
    body(doc,
         "Versi Flask 3.x yang digunakan mendukung Blueprint untuk memisahkan endpoint "
         "dashboard, cameras, dan servers. Flask-SQLAlchemy digunakan sebagai ORM untuk "
         "interaksi dengan SQLite, sedangkan Flask-APScheduler digunakan untuk penjadwalan "
         "background polling. Keunggulan Flask dalam konteks ini: bobot ringan, kemudahan "
         "membuat endpoint MJPEG streaming, dan dukungan streaming response untuk data real-time.")

    section_h(doc, "2.5", "Teknologi Frontend Web")
    body(doc,
         "Bootstrap 5.3 adalah framework CSS responsif yang menyediakan komponen UI siap "
         "pakai seperti card, table, navbar, modal, dan grid system 12-kolom. Dukungan "
         "dark mode bawaan melalui atribut data-bs-theme menjadi keunggulan Bootstrap 5.3 "
         "yang dimanfaatkan dalam penelitian ini (Bootstrap Team, 2023).")
    body(doc,
         "Leaflet.js adalah library JavaScript open-source untuk peta interaktif yang "
         "ringan (~42KB). Dalam dashboard ini, Leaflet menampilkan lokasi geografis kamera "
         "CCTV pada peta interaktif dengan fitur klik-untuk-tandai koordinat kamera baru "
         "menggunakan tile layer OpenStreetMap (Agafonkin, 2022). Font Awesome 6.5 "
         "digunakan sebagai library ikon vektor sehingga tampilan lebih informatif.")

    section_h(doc, "2.6", "Monitoring Server Storage")
    body(doc,
         "Dell iDRAC Redfish API: Redfish adalah spesifikasi RESTful API standar DMTF "
         "untuk manajemen infrastruktur server secara out-of-band menggunakan format JSON "
         "dan protokol HTTPS (DMTF, 2022). Endpoint yang digunakan: /redfish/v1/Systems "
         "untuk informasi sistem, /Chassis/Thermal untuk data suhu, /Storage untuk "
         "informasi disk, dan /Power untuk data power supply.")
    body(doc,
         "SNMP (Simple Network Management Protocol) adalah protokol standar Internet untuk "
         "manajemen perangkat jaringan pada UDP port 161. Untuk Pelco Endura NSM5200, "
         "OID yang digunakan berasal dari HOST-RESOURCES-MIB (hrStorageTable untuk storage) "
         "dan MIB-2 standar (sysDescr, sysName). Indikator kesehatan HDD yang dipantau: "
         "suhu drive, status kesehatan (OK/Warning/Critical), data SMART, total jam nyala, "
         "dan prediksi sisa umur pakai SSD.")

    section_h(doc, "2.7", "Sistem Basis Data SQLite")
    body(doc,
         "SQLite adalah RDBMS yang ringan, serverless, dan self-contained — menyimpan "
         "seluruh basis data dalam satu file tunggal. Hal ini menjadikannya sangat cocok "
         "untuk aplikasi skala menengah seperti dashboard monitoring yang tidak memerlukan "
         "konkurensi tinggi (Owens & Allen, 2010). Interaksi dengan SQLite dilakukan "
         "melalui Flask-SQLAlchemy (ORM) dengan empat entitas: Camera, Server, HDD, PSU.")

    section_h(doc, "2.8", "Kerangka Berpikir")
    body(doc,
         "Kerangka berpikir penelitian ini disusun berdasarkan identifikasi masalah, "
         "solusi yang diusulkan, dan hasil yang diharapkan. Gambar 2.1 menampilkan "
         "kerangka berpikir penelitian secara sistematis.")

    add_figure(doc, fc_path("kerangka_berpikir"),
               "Gambar 2.1  Kerangka Berpikir Penelitian", width_cm=9)


# ════════════════════════════════════════════════════════════════════════════
#  BAB III – METODOLOGI PENELITIAN
# ════════════════════════════════════════════════════════════════════════════
def build_bab3(doc):
    doc.add_page_break()
    chapter_title(doc, "BAB III\nMETODOLOGI PENELITIAN")

    section_h(doc, "3.1", "Jenis Penelitian")
    body(doc,
         "Penelitian ini termasuk dalam kategori penelitian rancang bangun (design and "
         "development research) atau penelitian terapan yang menghasilkan produk berupa "
         "perangkat lunak. Pendekatan yang digunakan adalah rekayasa perangkat lunak "
         "dengan metode pengembangan prototipe (prototyping model), yang memungkinkan "
         "evaluasi bertahap dan perbaikan iteratif berdasarkan hasil pengujian setiap "
         "iterasi (Borg & Gall, 1983). Selain itu, digunakan pula analisis kualitatif "
         "untuk mengidentifikasi kebutuhan sistem dan mengevaluasi hasil implementasi.")

    section_h(doc, "3.2", "Waktu dan Tempat Penelitian")
    body(doc,
         "Penelitian dilaksanakan selama ± 6 bulan (Januari – Juni 2026) di Politeknik "
         "Penerbangan Indonesia Curug, Jalan Raya PLP Curug, Kabupaten Tangerang, Banten. "
         "Pengujian sistem dilakukan di laboratorium komputer dan jaringan PPIC Curug "
         "menggunakan perangkat simulasi (demo mode) yang mengemulasi kamera Pelco "
         "dan server storage Pelco VX Storage.")

    section_h(doc, "3.3", "Teknik Pengumpulan Data")
    body(doc, "Data dikumpulkan melalui beberapa teknik berikut:")

    for judul_t, isi_t in [
        ("Studi Dokumentasi",
         "Mempelajari dokumentasi resmi ONVIF Forum, Dell Technologies (iDRAC Redfish API), "
         "Pelco, serta dokumentasi library Python yang digunakan (Flask, onvif-zeep, pysnmp, opencv)."),
        ("Observasi Langsung",
         "Pengamatan terhadap infrastruktur CCTV di laboratorium, termasuk respons kamera "
         "ONVIF terhadap probe dan cara iDRAC menyajikan data hardware melalui Redfish API."),
        ("Wawancara Tidak Terstruktur",
         "Diskusi dengan teknisi sistem keamanan dan dosen pembimbing untuk memahami kebutuhan "
         "dan permasalahan operasional sistem CCTV bandar udara."),
        ("Studi Literatur",
         "Mengkaji jurnal ilmiah, buku teks, prosiding, dan artikel teknis tentang CCTV, "
         "ONVIF, Flask, monitoring hardware, dan pengembangan dashboard berbasis web."),
        ("Eksperimen",
         "Percobaan langsung menghubungkan aplikasi ke perangkat kamera ONVIF dan server "
         "storage untuk memvalidasi fungsionalitas sistem."),
    ]:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.first_line_indent = Cm(1.25)
        r1 = p.add_run(judul_t + ": ")
        set_font(r1, size=12, bold=True)
        r2 = p.add_run(isi_t)
        set_font(r2, size=12)

    section_h(doc, "3.4", "Perancangan Sistem")
    body(doc,
         "Perancangan sistem menggunakan pendekatan berlapis (layered architecture) yang "
         "terdiri dari empat lapisan utama: presentasi (frontend), routing, layanan bisnis "
         "(service layer), dan data. Gambaran arsitektur keseluruhan sistem telah ditunjukkan "
         "pada Gambar 1.1 di Bab I. Pada bab ini diuraikan perancangan masing-masing "
         "komponen secara lebih rinci.")

    subsection_h(doc, "3.4.1", "Perancangan Basis Data")
    body(doc,
         "Basis data sistem terdiri dari empat tabel utama yang saling berelasi seperti "
         "ditampilkan pada Tabel 3.1 berikut.")

    add_table(doc,
        ["Tabel", "Kolom Utama", "Keterangan"],
        [
            ("camera",
             "id, name, ip_address, port, onvif_username, onvif_password, manufacturer, "
             "model, firmware, location_name, latitude, longitude, stream_uri, "
             "snapshot_uri, is_active, last_seen",
             "Data konfigurasi dan status kamera CCTV"),
            ("server",
             "id, name, ip_address, server_type, idrac_ip, idrac_username, idrac_password, "
             "snmp_community, system_model, power_state, health_rollup, inlet_temp, "
             "exhaust_temp, is_online, last_checked",
             "Data konfigurasi dan status server storage"),
            ("hdd",
             "id, server_id (FK), device_name, slot, model, serial, media_type, protocol, "
             "capacity_gb, used_gb, temperature_c, health_status, power_on_hours, "
             "reallocated_sectors, predicted_life_left",
             "Data setiap HDD/SSD pada server storage"),
            ("psu",
             "id, server_id (FK), name, model, health_status, power_watts, "
             "capacity_watts, last_checked",
             "Data setiap power supply unit pada server"),
        ],
        "Tabel 3.1  Skema Basis Data Sistem")

    subsection_h(doc, "3.4.2", "Perancangan Antarmuka Pengguna")
    body(doc,
         "Antarmuka pengguna dirancang dengan prinsip mobile-first responsive design "
         "menggunakan Bootstrap 5.3. Terdapat empat halaman utama: (1) Dashboard — "
         "statistik, peta kamera, dan status server; (2) Cameras List — daftar kamera "
         "dengan filter status aktif/tidak aktif; (3) Camera Detail — live stream MJPEG "
         "dan spesifikasi kamera; (4) Servers — kartu server storage dengan informasi "
         "kesehatan hardware (suhu, HDD, PSU).")
    body(doc,
         "Seluruh halaman menggunakan tata letak sidebar fixed (260px) di kiri untuk "
         "navigasi utama. Pada layar di bawah 768px, sidebar disembunyikan melalui "
         "tombol hamburger. Sistem mendukung dark mode dan light mode dengan preferensi "
         "tersimpan di localStorage browser.")

    subsection_h(doc, "3.4.3", "Alur Perancangan Dashboard")
    body(doc,
         "Proses perancangan dashboard monitoring CCTV dilakukan secara bertahap mulai "
         "dari analisis kebutuhan hingga dokumentasi sistem. Gambar 3.1 menampilkan "
         "flowchart alur perancangan dashboard secara keseluruhan.")

    add_figure(doc, fc_path("alur_perancangan"),
               "Gambar 3.1  Flowchart Alur Perancangan Dashboard", width_cm=10)

    subsection_h(doc, "3.4.4", "Perancangan Pemantauan Kamera")
    body(doc,
         "Sistem pemantauan kamera bekerja secara otomatis dengan memeriksa status setiap "
         "kamera yang terdaftar secara berkala. Gambar 3.2 menampilkan flowchart alur "
         "pemantauan status kamera.")

    add_figure(doc, fc_path("polling_kamera"),
               "Gambar 3.2  Flowchart Pemantauan Status Kamera", width_cm=10)

    subsection_h(doc, "3.4.5", "Perancangan Pemantauan Server Storage")
    body(doc,
         "Sistem pemantauan server storage menangani dua jenis server dengan cara yang "
         "berbeda sesuai karakteristik masing-masing. Gambar 3.3 menampilkan flowchart "
         "alur pemantauan server storage.")

    add_figure(doc, fc_path("polling_server"),
               "Gambar 3.3  Flowchart Pemantauan Server Storage", width_cm=13)

    section_h(doc, "3.5", "Diagram Alir Penelitian")
    body(doc,
         "Penelitian ini dilaksanakan mengikuti diagram alir sistematis seperti ditunjukkan "
         "pada Gambar 3.4. Proses dimulai dari identifikasi masalah dan studi literatur, "
         "kemudian perancangan sistem, implementasi, pengujian, evaluasi, hingga penulisan "
         "laporan akhir.")

    add_figure(doc, fc_path("alir_penelitian"),
               "Gambar 3.4  Diagram Alir Penelitian", width_cm=10)

    section_h(doc, "3.6", "Jadwal Kegiatan Penelitian")
    body(doc,
         "Penelitian direncanakan berlangsung selama 6 bulan (Januari – Juni 2026). "
         "Tabel 3.2 menampilkan rincian jadwal kegiatan penelitian per bulan.")

    add_table(doc,
        ["No.", "Kegiatan", "Jan", "Feb", "Mar", "Apr", "Mei", "Jun"],
        [
            ("1", "Studi literatur & identifikasi kebutuhan",      "✓", "✓", "",  "",  "",  ""),
            ("2", "Perancangan arsitektur & basis data",           "",  "✓", "✓", "",  "",  ""),
            ("3", "Implementasi backend Flask & integrasi ONVIF",  "",  "",  "✓", "✓", "",  ""),
            ("4", "Implementasi frontend & integrasi Redfish/SNMP","",  "",  "",  "✓", "✓", ""),
            ("5", "Pengujian fungsional & perbaikan",              "",  "",  "",  "",  "✓", "✓"),
            ("6", "Analisis hasil & penulisan laporan",            "",  "",  "",  "",  "✓", "✓"),
            ("7", "Seminar hasil & revisi laporan akhir",          "",  "",  "",  "",  "",  "✓"),
        ],
        "Tabel 3.2  Jadwal Kegiatan Penelitian")


# ════════════════════════════════════════════════════════════════════════════
#  DAFTAR PUSTAKA
# ════════════════════════════════════════════════════════════════════════════
def build_daftar_pustaka(doc):
    doc.add_page_break()
    p = doc.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(18)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    r = p.add_run("DAFTAR PUSTAKA")
    set_font(r, size=14, bold=True)

    refs = [
        "Agafonkin, V. (2022). Leaflet: An Open-Source JavaScript Library for "
        "Mobile-Friendly Interactive Maps. Leaflet.js Documentation. "
        "https://leafletjs.com (diakses 15 Januari 2026).",

        "Borg, W. R., & Gall, M. D. (1983). Educational Research: An Introduction "
        "(4th ed.). Longman.",

        "Bootstrap Team. (2023). Bootstrap 5.3 Documentation. "
        "https://getbootstrap.com/docs/5.3 (diakses 10 Januari 2026).",

        "Dell Technologies. (2023). Integrated Dell Remote Access Controller 9 "
        "User's Guide — Redfish API Guide. Dell Technical Documentation.",

        "DMTF. (2022). Redfish Specification Version 1.16.0. Distributed Management "
        "Task Force. https://www.dmtf.org/dsp/DSP0266.",

        "Grinberg, M. (2018). Flask Web Development: Developing Web Applications "
        "with Python (2nd ed.). O'Reilly Media.",

        "Huang, F. (2021). onvif-zeep: Python ONVIF Client Library. "
        "https://github.com/FalkTannhaeuser/python-onvif-zeep.",

        "ICAO. (2017). Annex 17 to the Convention on International Civil Aviation: "
        "Security — Safeguarding International Civil Aviation Against Acts of "
        "Unlawful Interference (10th ed.). International Civil Aviation Organization.",

        "Motorola Solutions. (2023). Pelco Camera Product Guide. "
        "https://www.pelco.com (diakses 20 Januari 2026).",

        "Norris, C., & Armstrong, G. (2019). The Maximum Surveillance Society: "
        "The Rise of CCTV (2nd ed.). Routledge.",

        "ONVIF. (2023). ONVIF Core Specification v23.06. "
        "https://www.onvif.org/specs/core/ONVIF-Core-Specification.pdf.",

        "Owens, M., & Allen, G. (2010). The Definitive Guide to SQLite (2nd ed.). Apress.",

        "Pelco. (2022). Pelco Sarix IMP/IME Series Camera Installation and "
        "Operation Manual. Pelco Technical Publications.",

        "Stallings, W. (2021). SNMP, SNMPv2, SNMPv3, and RMON 1 and 2 (3rd ed.). "
        "Addison-Wesley.",
    ]

    for ref in refs:
        p_r = doc.add_paragraph()
        p_r.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p_r.paragraph_format.space_before = Pt(0)
        p_r.paragraph_format.space_after = Pt(4)
        p_r.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        p_r.paragraph_format.left_indent = Cm(1.25)
        p_r.paragraph_format.first_line_indent = Cm(-1.25)   # hanging indent
        r = p_r.add_run(ref)
        set_font(r, size=12)


# ════════════════════════════════════════════════════════════════════════════
#  MAIN
# ════════════════════════════════════════════════════════════════════════════
def main():
    print("=" * 55)
    print("  Generator Proposal Tugas Akhir — PPIC Curug")
    print("=" * 55)

    generate_all_flowcharts()

    doc = Document()
    section = doc.sections[0]
    set_margins(section)
    section.page_height = Cm(29.7)
    section.page_width  = Cm(21.0)

    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)
    style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    style.paragraph_format.line_spacing = 1.5

    steps = [
        ("Halaman Judul",      build_cover),
        ("Halaman Persetujuan", build_approval),
        ("Daftar Isi",         build_daftar_isi),
        ("BAB I",              build_bab1),
        ("BAB II",             build_bab2),
        ("BAB III",            build_bab3),
        ("Daftar Pustaka",     build_daftar_pustaka),
    ]
    for i, (label, fn) in enumerate(steps, 1):
        print(f"  [{i}/{len(steps)}] {label} ...")
        fn(doc)

    doc.save(OUTPUT_PATH)
    print(f"\nDokumen berhasil disimpan ke:\n  {OUTPUT_PATH}")
    print("\nPlaceholder yang perlu diganti sebelum dikumpulkan:")
    print("  • MAHASISWA_NIM   — NIM asli")
    print("  • DOSEN_1, DOSEN_2 — nama dosen pembimbing")
    print("  • [LOGO PPIC CURUG] — sisipkan logo PPIC secara manual di Word")


if __name__ == "__main__":
    main()
